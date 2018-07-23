
# coding: utf-8

# ### 抬头式

# In[94]:



from docx import Document
from openpyxl import load_workbook
from docx import *
import time
import sys
import os


# ### 设置环境

# In[95]:


reload(sys)
sys.setdefaultencoding('utf8')


# ### 读取excel

# In[96]:


#把所有输入的当做是字符串
filename = raw_input("input filename:")
filenames = 'C:\\Users\\jordan\\Desktop\\FDC@FDCOAtable\\FDC@FDC_v2\\{}.xlsx'.format(filename)


# In[97]:


wb = load_workbook(filename=filenames)
#wb = load_workbook(filename=r'C:\Users\jordan\Desktop\BOND_BASICINFO.xlsx')
#title = str(wb.get_sheet_names()[0])
#获取当前正在使用的sheet
sheet = wb.active 
title = u'' + str(wb.get_sheet_names()[0])
title


# In[98]:


### 加title

document = Document()
document.add_heading(title, level=1)


# In[99]:


table = document.add_table(rows=2, cols=2)

#获取第一行的單元格對象
#col_cells = table.rows[0].cells


table.rows[0].cells[0].text = u'表名'
#table.rows[0].cells[1].text = 

table.rows[1].cells[0].text = u'表描述'
#table.rows[1].cells[1].text = 


# In[100]:


# 取第一张表
#获取所有的sheet
sheetnames = wb.get_sheet_names()
#取第一张sheet
ws = wb.get_sheet_by_name(sheetnames[0])
ws_rows = list(ws.rows)


# In[101]:


ws_rows_one = list(ws_rows[0])
ws_rows_one[0].value
type(ws_rows_one[0].value)
ws_rows_one[0].value


# ### 加A3和A4

# In[102]:


#取A3
ws_rows_three = list(ws_rows[2])
ws_rows_three = ws_rows_three[0]
ws_rows_three = ws_rows_three.value
ws_rows_three = ws_rows_three.split(":")[1]
ws_rows_three

table.rows[0].cells[1].text = ws_rows_three

#取A4
ws_rows_four = list(ws_rows[3])
ws_rows_four = ws_rows_four[0]
ws_rows_four = ws_rows_four.value
ws_rows_four = ws_rows_four.split(":")[1]
ws_rows_four

table.rows[1].cells[1].text = ws_rows_four


# ### 建立底层的TABLE

# In[103]:


#获得当前最大rows
rows = sheet.max_row
#获得当前最大cols
cols = sheet.max_column

#先建立标题table
title_table = document.add_table(rows=1, cols=cols)
#获取第一行的單元格對象
title_cells = title_table.rows[0].cells
title_cells[0].text = u'序號'
title_cells[1].text = u'主鍵'
title_cells[2].text = u'代碼'
title_cells[3].text = u'名稱'
title_cells[4].text = u'類型'
title_cells[5].text = u'可空'
title_cells[6].text = u'備註'


# In[104]:


#在建立底部table
bottom_table = document.add_table(rows=rows-1, cols=cols)
for i in range(0,rows-5):
    bottom_cells = bottom_table.rows[i].cells
    bottom_cells[0].text = u'' + str(ws_rows[i+5][0].value)
    bottom_cells[1].text = u'' + str(ws_rows[i+5][1].value)
    bottom_cells[2].text = u'' + str(ws_rows[i+5][2].value)
    bottom_cells[3].text = u'' + str(ws_rows[i+5][3].value)
    bottom_cells[4].text = u'' + str(ws_rows[i+5][4].value)
    bottom_cells[5].text = u'' + str(ws_rows[i+5][5].value)
    bottom_cells[6].text = u'' + str(ws_rows[i+5][6].value)


# In[105]:


document.save(r'C:\Users\jordan\Desktop\FDC_WORD\{}.docx'.format(filename))


# ### docx 转html

# In[107]:



# In[10]:


### 党名自动转成test03.html


# In[86]:





# In[116]:




